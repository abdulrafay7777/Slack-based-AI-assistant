"""
DOCX generation using python-docx.

Generates a formatted proposal document from the Writer Agent's section dict.
Each section is rendered with proper headings, body text, and spacing.
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(os.getenv("DOCX_OUTPUT_DIR", "outputs"))



BRAND_COLOR = RGBColor(0x1F, 0x49, 0x7D)   # deep navy


def _set_heading_style(paragraph, level: int = 1) -> None:
    """Apply brand heading style to a paragraph."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.color.rgb = BRAND_COLOR
    run.font.bold = True
    run.font.size = Pt(18) if level == 1 else Pt(14)
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line using paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)


def _add_cover_section(doc: Document, text: str, client_data: Dict[str, Any]) -> None:
    """Render the cover section with extra prominence."""
    company = (client_data.get("company_name") or "Client").title()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"Proposal for {company}")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BRAND_COLOR

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_horizontal_rule(doc)
    doc.add_paragraph(text)


def _add_section(doc: Document, title: str, text: str) -> None:
    """Render a standard proposal section with a heading and body paragraphs."""
    heading = doc.add_heading(title, level=1)
    _set_heading_style(heading)
    _add_horizontal_rule(doc)

    # Split on double newlines to preserve paragraph breaks
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)


def _set_document_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


async def build_docx(
    sections: list[Dict[str, str]],
    client_data: Dict[str, Any],
    session_id: str,
    revision: int = 0,
) -> str:
    """
    Build a DOCX proposal from the Writer Agent's dynamic sections list.

    Args:
        sections:    List of dicts with 'id', 'title', and 'content'.
        client_data: Structured client info for the cover page.
        session_id:  Used to generate a unique filename.
        revision:    Revision counter appended to filename.

    Returns:
        Absolute path to the generated DOCX file.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_document_margins(doc)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Render sections in order
    for i, section_data in enumerate(sections):
        section_id = section_data.get("id", "")
        section_title = section_data.get("title", "")
        text = section_data.get("content", "").strip()
        
        if not text:
            continue

        if section_id == "cover":
            _add_cover_section(doc, text, client_data)
        else:
            if i > 1:   # add page break before every section except the first one after the cover
                doc.add_page_break()
            _add_section(doc, section_title, text)

    # Footer with session info
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Confidential · Proposal ID: {session_id[:8]} · Rev {revision}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    suffix = f"_rev{revision}" if revision > 0 else ""
    filename = f"proposal_{session_id[:8]}{suffix}.docx"
    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))

    return str(out_path.resolve())
